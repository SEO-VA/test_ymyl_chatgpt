#!/usr/bin/env python3
"""
Content Processing Web Application with AI Compliance Analysis

This script combines the robust content extraction and chunking with
parallel AI-powered YMYL compliance analysis capabilities.
"""

import streamlit as st
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, WebDriverException
import json
import time
import html
from datetime import datetime
import pytz
import platform
import logging
import asyncio
import aiohttp
from openai import OpenAI
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import markdown
import re

# --- Logging Configuration ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Streamlit Page Configuration ---
st.set_page_config(
    page_title="YMYL Audit Tool",
    page_icon="🚀",
    layout="wide",
)

# --- AI Processing Configuration ---
ANALYZER_ASSISTANT_ID = "asst_WzODK9EapCaZoYkshT6x9xEH"

# --- Component 1: Updated Content Extractor ---
class ContentExtractor:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def extract_content(self, url):
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            
            content_parts = []
            
            # 1. Extract H1 (anywhere on page)
            h1 = soup.find('h1')
            if h1:
                text = h1.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"H1: {text}")
            
            # 2. Extract Subtitle (anywhere on page)
            subtitle = soup.find('span', class_=['sub-title', 'd-block'])
            if subtitle:
                text = subtitle.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"SUBTITLE: {text}")
            
            # 3. Extract Lead (anywhere on page)
            lead = soup.find('p', class_='lead')
            if lead:
                text = lead.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"LEAD: {text}")
            
            # 4. Extract Article content
            article = soup.find('article')
            if article:
                # Remove tab-content sections before processing
                for tab_content in article.find_all('div', class_='tab-content'):
                    tab_content.decompose()
                
                # Process all elements in document order within article
                for element in article.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'p']):
                    text = element.get_text(separator='\n', strip=True)
                    if not text:
                        continue
                    
                    # Check element type and add appropriate prefix
                    if element.name == 'h1':
                        content_parts.append(f"H1: {text}")
                    elif element.name == 'h2':
                        content_parts.append(f"H2: {text}")
                    elif element.name == 'h3':
                        content_parts.append(f"H3: {text}")
                    elif element.name == 'h4':
                        content_parts.append(f"H4: {text}")
                    elif element.name == 'h5':
                        content_parts.append(f"H5: {text}")
                    elif element.name == 'h6':
                        content_parts.append(f"H6: {text}")
                    elif element.name == 'span' and 'sub-title' in element.get('class', []) and 'd-block' in element.get('class', []):
                        content_parts.append(f"SUBTITLE: {text}")
                    elif element.name == 'p' and 'lead' in element.get('class', []):
                        content_parts.append(f"LEAD: {text}")
                    elif element.name == 'p':
                        content_parts.append(f"CONTENT: {text}")
            
            # 5. Extract FAQ section
            faq_section = soup.find('section', attrs={'data-qa': 'templateFAQ'})
            if faq_section:
                text = faq_section.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"FAQ: {text}")
            
            # 6. Extract Author section
            author_section = soup.find('section', attrs={'data-qa': 'templateAuthorCard'})
            if author_section:
                text = author_section.get_text(separator='\n', strip=True)
                if text:
                    content_parts.append(f"AUTHOR: {text}")
            
            # Join with double newlines to preserve spacing
            final_content = '\n\n'.join(content_parts)
            return True, final_content, None
            
        except requests.RequestException as e:
            return False, None, f"Error fetching URL: {e}"
        except Exception as e:
            return False, None, f"Error processing content: {e}"

# --- Component 2: The Final, Upgraded Chunk Processor ---
class ChunkProcessor:
    def __init__(self, log_callback=None):
        self.driver = None
        self.log = log_callback if log_callback else logger.info

    def _setup_driver(self):
        self.log("Initializing browser with enhanced stability & permissions...")
        chrome_options = Options()
        chrome_options.add_argument('--headless=new')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_experimental_option("prefs", {"profile.default_content_setting_values.clipboard": 1})
        try:
            self.driver = webdriver.Chrome(options=chrome_options)
            self.log("✅ Browser initialized successfully.")
            return True
        except WebDriverException as e:
            self.log(f"❌ WebDriver Initialization Failed: {e}")
            return False

    def _extract_json_from_button(self):
        try:
            wait = WebDriverWait(self.driver, 180)
            h3_xpath = "//h3[text()='Raw JSON Output']"
            self.log("🔄 Waiting for results section to appear...")
            wait.until(EC.presence_of_element_located((By.XPATH, h3_xpath)))
            self.log("✅ Results section is visible.")
            button_selector = "button[data-testid='stCodeCopyButton']"
            self.log("...Waiting for the copy button...")
            copy_button = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, button_selector)))
            self.log("✅ Found the copy button element.")
            self.log("...Polling button's attribute for completeness...")
            timeout = time.time() + 10
            final_content = ""
            while time.time() < timeout:
                raw_content = copy_button.get_attribute('data-clipboard-text')
                if raw_content and raw_content.strip().startswith('{') and raw_content.strip().endswith('}'):
                    final_content = raw_content; break
                time.sleep(0.2)
            if not final_content: self.log("❌ Timed out polling the attribute."); return None
            self.log("...Decoding HTML entities...")
            decoded_content = html.unescape(final_content)
            self.log(f"✅ Extraction complete. Retrieved {len(decoded_content):,} characters.")
            return decoded_content
        except Exception as e:
            self.log(f"❌ An error occurred during the final JSON extraction phase: {e}")
            return None

    def process_content(self, content):
        if not self._setup_driver():
            return False, None, "Failed to initialize browser."
        try:
            self.log(f"Navigating to `chunk.dejan.ai`...")
            self.driver.get("https://chunk.dejan.ai/")
            wait = WebDriverWait(self.driver, 30)
            self.log("Using JavaScript to copy full text to browser's clipboard...")
            self.driver.execute_script("navigator.clipboard.writeText(arguments[0]);", content)
            self.log("Locating text area and clearing it...")
            textarea_selector = (By.CSS_SELECTOR, 'textarea[aria-label="Text to chunk:"]')
            input_field = wait.until(EC.element_to_be_clickable(textarea_selector))
            input_field.clear()
            self.log("Simulating a 'Paste' (Ctrl+V) command...")
            modifier_key = Keys.COMMAND if platform.system() == "Darwin" else Keys.CONTROL
            input_field.send_keys(modifier_key, "v")
            self.log("Clicking submit button...")
            submit_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '[data-testid="stBaseButton-secondary"]')))
            submit_button.click()
            json_output = self._extract_json_from_button()
            if json_output:
                return True, json_output, None
            else:
                return False, None, "Failed to extract JSON from the results page."
        except Exception as e:
            return False, None, f"An unexpected error occurred during processing: {e}"
        finally:
            self.cleanup()
            
    def cleanup(self):
        if self.driver:
            self.log("Cleaning up and closing browser instance.")
            self.driver.quit()
            self.log("✅ Browser closed.")

# --- Component 3: AI Processing Functions ---

def convert_to_html(markdown_content):
    """Convert markdown report to styled HTML"""
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(markdown_content, extensions=['tables', 'toc'])
        
        # Add professional CSS styling
        css_style = """
        <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 10px;
            margin-top: 30px;
        }
        h3 {
            color: #34495e;
            margin-top: 25px;
        }
        .severity-critical { color: #e74c3c; font-weight: bold; }
        .severity-high { color: #e67e22; font-weight: bold; }
        .severity-medium { color: #f39c12; font-weight: bold; }
        .severity-low { color: #3498db; font-weight: bold; }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        .processing-summary {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }
        code {
            background-color: #f1f2f6;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }
        blockquote {
            border-left: 4px solid #bdc3c7;
            margin: 0;
            padding-left: 15px;
            color: #7f8c8d;
        }
        </style>
        """
        
        # Enhance severity indicators
        html_content = html_content.replace('🔴', '<span class="severity-critical">🔴</span>')
        html_content = html_content.replace('🟠', '<span class="severity-high">🟠</span>')
        html_content = html_content.replace('🟡', '<span class="severity-medium">🟡</span>')
        html_content = html_content.replace('🔵', '<span class="severity-low">🔵</span>')
        
        # Wrap processing summary
        html_content = re.sub(
            r'## Processing Summary(.*?)(?=##|$)', 
            r'<div class="processing-summary"><h2>Processing Summary</h2>\1</div>', 
            html_content, 
            flags=re.DOTALL
        )
        
        # Complete HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>YMYL Compliance Audit Report</title>
            {css_style}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        return full_html.encode('utf-8')
    except Exception as e:
        logger.error(f"HTML conversion error: {e}")
        return f"<html><body><h1>Export Error</h1><p>Failed to convert report: {e}</p></body></html>".encode('utf-8')

def convert_to_word(markdown_content):
    """Convert markdown report to Word document"""
    try:
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        
        # Create custom heading styles
        if 'Report Title' not in styles:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse markdown content
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.style = 'Report Title'
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('- '):
                # Bullet points
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('---'):
                # Horizontal rule (skip)
                continue
            elif '🔴' in line or '🟠' in line or '🟡' in line or '🔵' in line:
                # Severity indicators - make them stand out
                p = doc.add_paragraph(line)
                if '🔴' in line:
                    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)  # Red
                elif '🟠' in line:
                    p.runs[0].font.color.rgb = RGBColor(230, 126, 34)  # Orange
                elif '🟡' in line:
                    p.runs[0].font.color.rgb = RGBColor(243, 156, 18)  # Yellow/Gold
                elif '🔵' in line:
                    p.runs[0].font.color.rgb = RGBColor(52, 152, 219)  # Blue
                p.runs[0].font.bold = True
            else:
                # Regular paragraph
                if line:
                    doc.add_paragraph(line)
        
        # Save to memory
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Word conversion error: {e}")
        # Return simple document with error message
        doc = Document()
        doc.add_heading('Export Error', 0)
        doc.add_paragraph(f'Failed to convert report: {e}')
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

def convert_to_pdf(markdown_content):
    """Convert markdown report to PDF document"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        critical_style = ParagraphStyle(
            'Critical',
            parent=styles['Normal'],
            textColor=colors.red,
            fontSize=10,
            fontName='Helvetica-Bold'
        )
        
        high_style = ParagraphStyle(
            'High',
            parent=styles['Normal'],
            textColor=colors.orange,
            fontSize=10,
            fontName='Helvetica-Bold'
        )
        
        # Build story
        story = []
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                story.append(Spacer(1, 12))
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['Normal']))
            elif line.startswith('---'):
                story.append(Spacer(1, 12))
            elif '🔴' in line:
                story.append(Paragraph(line, critical_style))
            elif '🟠' in line:
                story.append(Paragraph(line, high_style))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            else:
                if line:
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        # Return simple error PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = [
            Paragraph("Export Error", getSampleStyleSheet()['Title']),
            Spacer(1, 12),
            Paragraph(f"Failed to convert report: {e}", getSampleStyleSheet()['Normal'])
        ]
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

def create_export_options(report_content):
    """Create multiple export format options"""
    return {
        'html': convert_to_html(report_content),
        'docx': convert_to_word(report_content),
        'pdf': convert_to_pdf(report_content),
        'markdown': report_content.encode('utf-8')
    }

def extract_big_chunks(json_data):
    """Extract and format big chunks for AI processing."""
    try:
        big_chunks = json_data.get('big_chunks', [])
        chunks = []
        
        for chunk in big_chunks:
            chunk_index = chunk.get('big_chunk_index', len(chunks) + 1)
            small_chunks = chunk.get('small_chunks', [])
            
            # Join small chunks with newlines
            joined_text = '\n'.join(small_chunks)
            
            chunks.append({
                "index": chunk_index,
                "text": joined_text,
                "count": len(small_chunks)
            })
        
        return chunks
    except Exception as e:
        logger.error(f"Error extracting big chunks: {e}")
        return []

async def call_assistant(api_key, assistant_id, content, chunk_index):
    """Call OpenAI Assistant API for chunk analysis."""
    try:
        client = OpenAI(api_key=api_key)
        
        # Create thread
        thread = client.beta.threads.create()
        logger.info(f"Thread created: {thread.id}")
        
        # Add message
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=content
        )
        
        # Run assistant
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )
        
        # Poll for completion
        while run.status in ['queued', 'in_progress']:
            await asyncio.sleep(1)
            run = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
        
        if run.status == 'completed':
            # Get response
            messages = client.beta.threads.messages.list(thread_id=thread.id)
            response_content = messages.data[0].content[0].text.value
            
            logger.info(f"Assistant call completed successfully for chunk {chunk_index}")
            return {
                "success": True,
                "content": response_content,
                "chunk_index": chunk_index
            }
        else:
            logger.error(f"Assistant run failed with status: {run.status}")
            return {
                "success": False,
                "error": f"Assistant run failed: {run.status}",
                "chunk_index": chunk_index
            }
            
    except Exception as e:
        logger.error(f"Error calling assistant for chunk {chunk_index}: {e}")
        return {
            "success": False,
            "error": str(e),
            "chunk_index": chunk_index
        }

async def process_chunks_parallel(chunks, api_key):
    """Process all chunks in parallel using OpenAI Assistant."""
    try:
        tasks = []
        for chunk in chunks:
            task = call_assistant(
                api_key=api_key,
                assistant_id=ANALYZER_ASSISTANT_ID,
                content=chunk["text"],
                chunk_index=chunk["index"]
            )
            tasks.append(task)
        
        # Execute all tasks simultaneously
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Process results and handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({
                    "success": False,
                    "error": str(result),
                    "chunk_index": chunks[i]["index"]
                })
            else:
                processed_results.append(result)
        
        return processed_results
        
    except Exception as e:
        logger.error(f"Error in parallel processing: {e}")
        return [{"success": False, "error": str(e), "chunk_index": 0}]

def create_final_report_simple(analysis_results):
    """Create final report by simple concatenation."""
    try:
        report_parts = []
        
        # Add header
        audit_date = datetime.now().strftime("%Y-%m-%d")
        header = f"""# YMYL Compliance Audit Report

**Audit Date:** {audit_date}
**Content Type:** Online Casino/Gambling  
**Analysis Method:** Section-by-section E-E-A-T compliance review

---

"""
        report_parts.append(header)
        
        # Add successful analyses
        successful_count = 0
        error_count = 0
        
        for result in analysis_results:
            if result.get("success"):
                report_parts.append(result["content"])
                report_parts.append("\n---\n")
                successful_count += 1
            else:
                error_count += 1
                error_section = f"""
# Analysis Error for Chunk {result.get('chunk_index', 'Unknown')}

❌ **Processing Failed**
Error: {result.get('error', 'Unknown error')}

---
"""
                report_parts.append(error_section)
        
        # Add processing summary
        total_sections = successful_count + error_count
        summary = f"""
## Processing Summary
**✅ Sections Successfully Analyzed:** {successful_count}
**❌ Sections with Analysis Errors:** {error_count}  
**📊 Total Sections:** {total_sections}

---
*Report generated by AI-powered YMYL compliance analysis system*
"""
        report_parts.append(summary)
        
        return ''.join(report_parts)
        
    except Exception as e:
        logger.error(f"Error creating final report: {e}")
        return f"Error generating report: {e}"

# --- Main Workflow Function ---
def process_url_workflow_with_logging(url, log_callback=None):
    result = {'success': False, 'url': url, 'extracted_content': None, 'json_output': None, 'error': None}
    
    def log(message):
        if log_callback: log_callback(message)
        logger.info(message)
        
    try:
        log("🚀 Initializing content extractor...")
        extractor = ContentExtractor()
        log(f"🔍 Fetching and extracting content from: {url}")
        success, content, error = extractor.extract_content(url)
        if not success:
            result['error'] = f"Content extraction failed: {error}"; return result
        result['extracted_content'] = content
        log(f"✅ Content extracted: {len(content):,} characters")

        log("🤖 Initializing chunk processor...")
        processor = ChunkProcessor(log_callback=log)
        success, json_output, error = processor.process_content(content)
        if not success:
            result['error'] = f"Chunk processing failed: {error}"; return result
        
        result['json_output'] = json_output
        result['success'] = True
        log("🎉 Workflow Complete!")
        return result
    except Exception as e:
        log(f"💥 An unexpected error occurred in the workflow: {str(e)}")
        result['error'] = f"An unexpected workflow error occurred: {str(e)}"
        return result

async def process_ai_analysis(json_output, api_key, log_callback=None):
    """Process AI compliance analysis on chunked content."""
    def log(message):
        if log_callback: log_callback(message)
        logger.info(message)
    
    try:
        # Parse JSON and extract chunks
        log("📊 Parsing JSON and extracting chunks...")
        json_data = json.loads(json_output)
        chunks = extract_big_chunks(json_data)
        
        if not chunks:
            return False, "No chunks found in JSON data", None
            
        log(f"🚀 Starting parallel analysis of {len(chunks)} chunks...")
        log(f"- Analyzer: {ANALYZER_ASSISTANT_ID}")
        log("- Report Maker: Simple Concatenation (No AI)")
        log("- API Key Status: ✅ Loaded")
        
        # Process chunks in parallel
        analysis_results = await process_chunks_parallel(chunks, api_key)
        
        # Create final report
        log("📝 Assembling final report...")
        final_report = create_final_report_simple(analysis_results)
        
        log("🎉 AI Analysis Complete!")
        return True, final_report, analysis_results
        
    except json.JSONDecodeError as e:
        error_msg = f"Invalid JSON format: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None
    except Exception as e:
        error_msg = f"AI analysis error: {e}"
        log(f"❌ {error_msg}")
        return False, error_msg, None

# --- Streamlit UI ---
def main():
    st.title("🕵 YMYL Audit Tool")
    st.markdown("**Automatically extract content from websites, generate JSON chunks, and perform YMYL compliance analysis**")

    # Sidebar configuration
    debug_mode = st.sidebar.checkbox("🐛 Debug Mode", value=True, help="Show detailed processing logs")

    # API Key configuration
    st.sidebar.markdown("### 🔑 AI Analysis Configuration")
    try:
        api_key = st.secrets["openai_api_key"]
        st.sidebar.success("✅ API Key loaded from secrets")
    except Exception:
        api_key = st.sidebar.text_input(
            "OpenAI API Key:",
            type="password",
            help="Enter your OpenAI API key for AI analysis"
        )
        if api_key:
            st.sidebar.success("✅ API Key provided")
        else:
            st.sidebar.warning("⚠️ API Key needed for AI analysis")

    st.markdown("---")
    col1, col2 = st.columns([2, 1])

    with col1:
        url = st.text_input("Enter the URL to process:", help="Include http:// or https://")
        if st.button("🚀 Process URL", type="primary", use_container_width=True):
            if not url:
                st.error("Please enter a URL to process")
                return

            # clear old state
            for key in ("latest_result", "ai_analysis_result"):
                st.session_state.pop(key, None)

            if debug_mode:
                # Detailed logging
                log_placeholder = st.empty()
                log_lines = []
                def log_callback(msg):
                    now = datetime.now(pytz.timezone("Europe/Malta"))
                    log_lines.append(f"`{now.strftime('%H:%M:%S')}`: {msg}")
                    log_placeholder.info("\n".join(log_lines))

                result = process_url_workflow_with_logging(url, log_callback)
                st.session_state["latest_result"] = result
                if result["success"]:
                    st.success("Processing completed successfully!")
                else:
                    st.error(f"Error: {result['error']}")
            else:
                # Simple milestones
                log_area = st.empty()
                milestones = []
                def simple_log(text):
                    milestones.append(f"- {text}")
                    log_area.markdown("\n".join(milestones))

                simple_log("Extracting content")
                extractor = ContentExtractor()
                ok, content, err = extractor.extract_content(url)
                if not ok:
                    st.error(f"Error: {err}")
                    return

                simple_log("Sending content to Chunk Norris")
                processor = ChunkProcessor()

                with st.status("You are not waiting, Chunk Norris is waiting for you"):
                    ok, json_out, err = processor.process_content(content)

                simple_log("Chunking done!")
                st.success("Chunking done!")

                st.session_state["latest_result"] = {
                    "success": ok,
                    "url": url,
                    "extracted_content": content if ok else None,
                    "json_output": json_out if ok else None,
                    "error": err
                }

    with col2:
        st.subheader("ℹ️ How it works")
        st.markdown("""
1. **Extract**: Extract the content.
2. **Chunk**: Send extracted text to Chunk Norris.
3. **YMYL Analysis**: YMYL audit of the content with AI.
4. **Done**: Output complete report.
""")
        st.info("💡 **New**: AI-powered YMYL compliance analysis available!")

    # Results Display
    if 'latest_result' in st.session_state and st.session_state['latest_result'].get('success'):
        result = st.session_state['latest_result']
        st.markdown("---")
        st.subheader("📊 Results")
        
        # AI Analysis Button
        if api_key and st.button("🤖 Process with AI Compliance Analysis", type="secondary", use_container_width=True):
            try:
                # Parse JSON and extract chunks first
                json_data = json.loads(result['json_output'])
                chunks = extract_big_chunks(json_data)
                
                if not chunks:
                    st.error("No chunks found in JSON data")
                    return
                
                # Enhanced Processing Logs Section
                st.subheader("🔍 Processing Logs")
                log_container = st.container()
                
                with log_container:
                    st.info(f"🚀 Starting parallel analysis of {len(chunks)} chunks...")
                    st.write("**Assistant IDs:**")
                    st.write(f"- Analyzer: `{ANALYZER_ASSISTANT_ID}`")
                    st.write("- Report Maker: `Simple Concatenation (No AI)`")
                    st.write(f"**API Key Status:** {'✅ Loaded' if api_key.startswith('sk-') else '❌ Invalid'}")
                    st.write("**Chunk Details:**")
                    for chunk in chunks:
                        st.write(f"- Chunk {chunk['index']}: {len(chunk['text']):,} characters")
                
                # Progress tracking
                total_chunks = len(chunks)
                progress_bar = st.progress(0)
                status_container = st.empty()
                
                # Start processing with timing
                start_time = time.time()
                
                with st.spinner("🤖 Running parallel analysis..."):
                    # Run AI analysis
                    success, ai_result, analysis_details = asyncio.run(process_ai_analysis(
                        result['json_output'], 
                        api_key, 
                        None  # Disable callback since we have enhanced UI
                    ))
                
                # Update progress
                progress_bar.progress(1.0)
                processing_time = time.time() - start_time
                
                # Display processing summary
                if success and analysis_details:
                    successful_analyses = [r for r in analysis_details if r.get("success")]
                    failed_analyses = [r for r in analysis_details if not r.get("success")]
                    
                    with status_container.container():
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Total Chunks", total_chunks)
                        with col2:
                            st.metric("Successful", len(successful_analyses), 
                                     delta=len(successful_analyses) if len(successful_analyses) == total_chunks else None)
                        with col3:
                            st.metric("Failed", len(failed_analyses), 
                                     delta=f"-{len(failed_analyses)}" if len(failed_analyses) > 0 else None)
                        
                        st.success(f"✅ Parallel analysis completed in {processing_time:.2f} seconds")
                    
                    # Store results
                    st.session_state['ai_analysis_result'] = {
                        'success': True,
                        'report': ai_result,
                        'details': analysis_details,
                        'processing_time': processing_time,
                        'total_chunks': total_chunks,
                        'successful_count': len(successful_analyses),
                        'failed_count': len(failed_analyses)
                    }
                    
                else:
                    st.session_state['ai_analysis_result'] = {
                        'success': False,
                        'error': ai_result if not success else 'Unknown error occurred'
                    }
                    st.error(f"❌ AI analysis failed: {ai_result if not success else 'Unknown error'}")
                    
            except json.JSONDecodeError as e:
                st.error(f"❌ Invalid JSON format: {str(e)}")
            except Exception as e:
                st.error(f"❌ An error occurred during AI analysis: {str(e)}")
                logger.error(f"AI analysis error: {str(e)}")

        # Tabs for results
        if 'ai_analysis_result' in st.session_state and st.session_state['ai_analysis_result'].get('success'):
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["🎯 AI Compliance Report", "📊 Individual Analyses", "🔧 JSON Output", "📄 Extracted Content", "📈 Summary"])
            
            with tab1:
                st.markdown("### YMYL Compliance Analysis Report")
                ai_report = st.session_state['ai_analysis_result']['report']
                
                # Enhanced Export Options
                st.markdown("#### 📋 Copy Report")
                st.code(ai_report, language='markdown')
                
                # Multiple Export Format Options
                st.markdown("#### 📄 Download Formats")
                st.markdown("Choose your preferred format for professional use:")
                
                # Create export data
                try:
                    export_formats = create_export_options(ai_report)
                    timestamp = int(time.time())
                    
                    # Create download buttons in columns
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.download_button(
                            label="📝 Markdown",
                            data=export_formats['markdown'],
                            file_name=f"ymyl_compliance_report_{timestamp}.md",
                            mime="text/markdown",
                            help="Original markdown format - perfect for copying to other platforms"
                        )
                    
                    with col2:
                        st.download_button(
                            label="🌐 HTML",
                            data=export_formats['html'],
                            file_name=f"ymyl_compliance_report_{timestamp}.html",
                            mime="text/html",
                            help="Styled HTML document - opens in any web browser"
                        )
                    
                    with col3:
                        st.download_button(
                            label="📄 Word",
                            data=export_formats['docx'],
                            file_name=f"ymyl_compliance_report_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Microsoft Word document - ready for editing and sharing"
                        )
                    
                    with col4:
                        st.download_button(
                            label="📋 PDF",
                            data=export_formats['pdf'],
                            file_name=f"ymyl_compliance_report_{timestamp}.pdf",
                            mime="application/pdf",
                            help="Professional PDF document - perfect for presentations and archival"
                        )
                    
                    st.info("""
                    💡 **Format Guide:**
                    - **Markdown**: Best for developers and copy-pasting to other platforms
                    - **HTML**: Opens in web browsers, styled and formatted
                    - **Word**: Professional business format, editable and shareable
                    - **PDF**: Final presentation format, preserves formatting across devices
                    """)
                    
                except Exception as e:
                    st.error(f"Error creating export formats: {e}")
                    # Fallback to basic markdown download
                    st.download_button(
                        label="💾 Download Report (Markdown)",
                        data=ai_report,
                        file_name=f"ymyl_compliance_report_{timestamp}.md",
                        mime="text/markdown"
                    )
                
                with st.expander("📖 View Formatted Report"):
                    st.markdown(ai_report)
            
            with tab2:
                st.markdown("### Individual Chunk Analysis Results")
                analysis_details = st.session_state['ai_analysis_result']['details']
                
                # Processing metrics at top
                ai_result = st.session_state['ai_analysis_result']
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Processing Time", f"{ai_result.get('processing_time', 0):.2f}s")
                with col2:
                    st.metric("Total Chunks", ai_result.get('total_chunks', 0))
                with col3:
                    st.metric("Successful", ai_result.get('successful_count', 0))
                with col4:
                    st.metric("Failed", ai_result.get('failed_count', 0))
                
                st.markdown("---")
                
                # Individual chunk results
                for detail in analysis_details:
                    chunk_idx = detail.get('chunk_index', 'Unknown')
                    if detail.get('success'):
                        with st.expander(f"✅ Chunk {chunk_idx} Analysis (Success)"):
                            st.markdown(detail['content'])
                            # Show additional metrics if available
                            if 'tokens_used' in detail:
                                st.caption(f"Tokens used: {detail['tokens_used']}")
                    else:
                        with st.expander(f"❌ Chunk {chunk_idx} Analysis (Failed)"):
                            st.error(f"Error: {detail.get('error', 'Unknown error')}")
            
            with tab3:
                st.code(result['json_output'], language='json')
                st.download_button(
                    label="💾 Download JSON",
                    data=result['json_output'],
                    file_name=f"chunks_{int(time.time())}.json",
                    mime="application/json"
                )
            
            with tab4:
                st.text_area("Raw extracted content:", value=result['extracted_content'], height=400)
            
            with tab5:
                st.subheader("Processing Summary")
                try:
                    json_data = json.loads(result['json_output'])
                    big_chunks = json_data.get('big_chunks', [])
                    total_small_chunks = sum(len(chunk.get('small_chunks', [])) for chunk in big_chunks)
                    
                    # Content extraction metrics
                    st.markdown("#### Content Extraction")
                    colA, colB, colC = st.columns(3)
                    colA.metric("Big Chunks", len(big_chunks))
                    colB.metric("Total Small Chunks", total_small_chunks)
                    colC.metric("Content Length", f"{len(result['extracted_content']):,} chars")
                    
                    # AI Analysis metrics
                    if 'ai_analysis_result' in st.session_state and st.session_state['ai_analysis_result'].get('success'):
                        st.markdown("#### AI Analysis Performance")
                        ai_result = st.session_state['ai_analysis_result']
                        analysis_details = ai_result['details']
                        successful_analyses = ai_result.get('successful_count', 0)
                        failed_analyses = ai_result.get('failed_count', 0)
                        processing_time = ai_result.get('processing_time', 0)
                        
                        colD, colE, colF, colG = st.columns(4)
                        colD.metric("Processing Time", f"{processing_time:.2f}s")
                        colE.metric("Successful Analyses", successful_analyses)
                        colF.metric("Failed Analyses", failed_analyses, 
                                   delta=f"-{failed_analyses}" if failed_analyses > 0 else None)
                        colG.metric("Success Rate", f"{(successful_analyses/(successful_analyses+failed_analyses)*100):.1f}%")
                        
                        # Performance insights
                        if processing_time > 0 and len(analysis_details) > 0:
                            avg_time_per_chunk = processing_time / len(analysis_details)
                            st.info(f"📊 **Performance**: Average {avg_time_per_chunk:.2f}s per chunk | Parallel efficiency achieved")
                        
                except (json.JSONDecodeError, TypeError):
                    st.warning("Could not parse JSON for statistics.")
                st.info(f"**Source URL**: {result['url']}")
        else:
            # Show original tabs without AI analysis
            tab1, tab2, tab3 = st.tabs(["🎯 JSON Output", "📄 Extracted Content", "📈 Summary"])
            
            with tab1:
                st.code(result['json_output'], language='json')
                st.download_button(
                    label="💾 Download JSON",
                    data=result['json_output'],
                    file_name=f"chunks_{int(time.time())}.json",
                    mime="application/json"
                )
            with tab2:
                st.text_area("Raw extracted content:", value=result['extracted_content'], height=400)
            with tab3:
                st.subheader("Processing Summary")
                try:
                    json_data = json.loads(result['json_output'])
                    big_chunks = json_data.get('big_chunks', [])
                    total_small_chunks = sum(len(chunk.get('small_chunks', [])) for chunk in big_chunks)
                    
                    colA, colB, colC = st.columns(3)
                    colA.metric("Big Chunks", len(big_chunks))
                    colB.metric("Total Small Chunks", total_small_chunks)
                    colC.metric("Content Length", f"{len(result['extracted_content']):,} chars")
                except (json.JSONDecodeError, TypeError):
                    st.warning("Could not parse JSON for statistics.")
                st.info(f"**Source URL**: {result['url']}")
        
        # Show API key reminder if not available
        if not api_key:
            st.info("💡 **Tip**: Add your OpenAI API key to enable AI compliance analysis!")

if __name__ == "__main__":
    main()
